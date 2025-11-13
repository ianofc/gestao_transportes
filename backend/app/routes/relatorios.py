import io
from flask import Blueprint, request, jsonify, send_file
from app import db
from app.models import CaixaDiario, Viagem, Usuario
from flask_jwt_extended import jwt_required
from dateutil import parser
from datetime import datetime

# Libs para PDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

# Libs para DOCX
from docx import Document
from docx.shared import Inches

bp = Blueprint('relatorios', __name__)

@bp.route('/caixa/<int:caixa_id>/pdf', methods=['GET'])
@jwt_required()
def relatorio_fecho_caixa_pdf(caixa_id):
    """
    Gera um relatório PDF para um fecho de caixa específico.
    """
    caixa = CaixaDiario.query.get_or_404(caixa_id)
    
    # Cria um buffer de bytes na memória para o PDF
    buffer = io.BytesIO()
    
    # Cria o canvas do PDF
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4 # (595.27, 841.89)
    
    # --- Conteúdo do PDF ---
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2.0, height - 2*cm, f"Relatório de Fecho de Caixa - ID: {caixa.id}")
    
    c.setFont("Helvetica", 12)
    y_pos = height - 3.5*cm
    
    # Função auxiliar para desenhar linhas
    def draw_line(label, value, y):
        c.drawString(2*cm, y, f"{label}:")
        c.drawString(7*cm, y, value)
        return y - 0.7*cm

    y_pos = draw_line("Bilheteiro", f"{caixa.bilheteiro.nome_completo if caixa.bilheteiro else 'N/A'}", y_pos)
    y_pos = draw_line("Status", caixa.status, y_pos)
    y_pos = draw_line("Abertura", f"{caixa.data_abertura.strftime('%d/%m/%Y %H:%M')}", y_pos)
    
    if caixa.data_fechamento:
        y_pos = draw_line("Fechamento", f"{caixa.data_fechamento.strftime('%d/%m/%Y %H:%M')}", y_pos)
    
    y_pos -= 0.5*cm # Espaçamento
    
    c.setFont("Helvetica-Bold", 12)
    y_pos = draw_line("Valores", "", y_pos)
    c.setFont("Helvetica", 12)

    y_pos = draw_line("Saldo Inicial", f"R$ {caixa.saldo_inicial:.2f}", y_pos)
    y_pos = draw_line("Vendas (Dinheiro)", f"R$ {caixa.total_vendas_dinheiro:.2f}", y_pos)
    y_pos = draw_line("Vendas (Pix)", f"R$ {caixa.total_vendas_pix:.2f}", y_pos)
    y_pos = draw_line("Vendas (Cartão)", f"R$ {caixa.total_vendas_cartao:.2f}", y_pos)
    
    y_pos -= 0.2*cm
    c.setStrokeColorRGB(0,0,0)
    c.line(2*cm, y_pos, width - 2*cm, y_pos) # Linha horizontal
    y_pos -= 0.7*cm
    
    c.setFont("Helvetica-Bold", 14)
    y_pos = draw_line("Total Geral em Vendas", f"R$ {caixa.total_geral_vendas:.2f}", y_pos)
    # --- Fim do Conteúdo ---

    c.showPage()
    c.save()
    
    # Retorna o buffer ao início
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f'relatorio_caixa_{caixa_id}.pdf',
        mimetype='application/pdf'
    )

@bp.route('/viagens/docx', methods=['GET'])
@jwt_required()
def relatorio_viagens_docx():
    """
    Gera um relatório DOCX das viagens (opcionalmente filtradas por data).
    Query Params: ?data_inicio=YYYY-MM-DD&data_fim=YYYY-MM-DD
    """
    data_inicio_str = request.args.get('data_inicio')
    data_fim_str = request.args.get('data_fim')
    
    query = Viagem.query.order_by(Viagem.data_partida_prevista.desc())
    
    periodo_str = "Período: Todas as viagens"
    
    try:
        if data_inicio_str:
            data_inicio = parser.parse(data_inicio_str).replace(hour=0, minute=0, second=0)
            query = query.filter(Viagem.data_partida_prevista >= data_inicio)
            periodo_str = f"Período de: {data_inicio.strftime('%d/%m/%Y')}"
            
        if data_fim_str:
            data_fim = parser.parse(data_fim_str).replace(hour=23, minute=59, second=59)
            query = query.filter(Viagem.data_partida_prevista <= data_fim)
            
            if data_inicio_str:
                periodo_str += f" até {data_fim.strftime('%d/%m/%Y')}"
            else:
                periodo_str = f"Período até: {data_fim.strftime('%d/%m/%Y')}"
                
    except Exception as e:
        return jsonify({"error": f"Formato de data inválido: {e}"}), 400

    viagens = query.all()
    
    # Cria documento DOCX na memória
    document = Document()
    buffer = io.BytesIO()

    # --- Conteúdo do DOCX ---
    document.add_heading('Relatório de Viagens', 0)
    document.add_paragraph(periodo_str)
    document.add_paragraph(f"Total de viagens encontradas: {len(viagens)}")
    
    if len(viagens) > 0:
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Rota (Origem-Destino)'
        hdr_cells[2].text = 'Partida Prevista'
        hdr_cells[3].text = 'Motorista'
        hdr_cells[4].text = 'Ônibus (Nº)'
        hdr_cells[5].text = 'Status'
        
        # Adiciona dados
        for v in viagens:
            row_cells = table.add_row().cells
            row_cells[0].text = str(v.id)
            row_cells[1].text = f"{v.rota.origem} - {v.rota.destino}" if v.rota else "N/A"
            row_cells[2].text = v.data_partida_prevista.strftime('%d/%m/%Y %H:%M')
            row_cells[3].text = v.motorista.nome_completo if v.motorista else "N/A"
            row_cells[4].text = v.onibus.numero_onibus if v.onibus else "N/A"
            row_cells[5].text = v.status
            
    # --- Fim do Conteúdo ---
    
    document.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name='relatorio_viagens.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )